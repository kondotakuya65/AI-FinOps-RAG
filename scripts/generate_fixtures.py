"""
Generate synthetic FinOps fixtures from a single in-code ground-truth source.

Usage (from repo root or backend/):
  python scripts/generate_fixtures.py
  # or
  cd backend && python -m scripts.generate_fixtures
"""

from __future__ import annotations

import json
import sys
from dataclasses import asdict, dataclass, field
from datetime import date
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

REPO_ROOT = Path(__file__).resolve().parents[1]
FIXTURES = REPO_ROOT / "fixtures"
INVOICES_DIR = FIXTURES / "invoices"
REPORTS_DIR = FIXTURES / "reports"
CONTRACTS_DIR = FIXTURES / "contracts"
EVALS_DIR = FIXTURES / "evals"


@dataclass
class LineItem:
    sku: str
    description: str
    qty: int
    unit_price: float

    @property
    def line_total(self) -> float:
        return round(self.qty * self.unit_price, 2)


@dataclass
class Invoice:
    invoice_id: str
    vendor: str
    invoice_date: str  # YYYY-MM-DD
    currency: str
    lines: list[LineItem]
    layout: str  # table_top | table_middle | table_bottom
    column_order: str  # qty_then_price | price_then_qty
    file_name: str = ""

    def __post_init__(self) -> None:
        if not self.file_name:
            self.file_name = f"{self.invoice_id.lower()}.pdf"

    @property
    def total_amount(self) -> float:
        return round(sum(line.line_total for line in self.lines), 2)

    @property
    def quarter(self) -> str:
        d = date.fromisoformat(self.invoice_date)
        q = (d.month - 1) // 3 + 1
        return f"{d.year}-Q{q}"


@dataclass
class ReportRow:
    sku: str
    vendor: str
    sold_qty: int
    received_qty: int
    return_rate: float


@dataclass
class MonthlyReport:
    period: str  # YYYY-MM
    rows: list[ReportRow]
    file_name: str = ""

    def __post_init__(self) -> None:
        if not self.file_name:
            self.file_name = f"product_report_{self.period.replace('-', '_')}.xlsx"


@dataclass
class Discrepancy:
    id: str
    sku: str
    invoice_id: str
    invoice_qty: int
    report_period: str
    report_received_qty: int
    severity: str = "quantity_mismatch"


@dataclass
class ContractPrice:
    sku: str
    unit_price: float


@dataclass
class Contract:
    vendor: str
    payment_terms: str
    effective_date: str
    prices: list[ContractPrice]
    file_name: str = "vendor_alpha_contract.docx"


@dataclass
class Corpus:
    vendors: list[str]
    invoices: list[Invoice]
    reports: list[MonthlyReport]
    discrepancies: list[Discrepancy]
    contract: Contract
    notes: list[str] = field(default_factory=list)


def build_corpus() -> Corpus:
    """Authoritative demo corpus — generators write files FROM this data."""
    invoices = [
        Invoice(
            invoice_id="INV-101",
            vendor="Alpha Supplies",
            invoice_date="2024-07-15",
            currency="USD",
            layout="table_top",
            column_order="qty_then_price",
            lines=[
                LineItem("SKU-1001", "Widget A", 100, 10.00),
                LineItem("SKU-1002", "Widget B", 50, 25.00),
            ],
        ),
        Invoice(
            invoice_id="INV-102",
            vendor="Alpha Supplies",
            invoice_date="2024-08-20",
            currency="USD",
            layout="table_middle",
            column_order="qty_then_price",
            lines=[
                LineItem("SKU-1001", "Widget A", 500, 10.00),
                LineItem("SKU-1003", "Widget C", 20, 21.511),  # 430.22
            ],
        ),
        Invoice(
            invoice_id="INV-103",
            vendor="Alpha Supplies",
            invoice_date="2024-09-10",
            currency="USD",
            layout="table_bottom",
            column_order="price_then_qty",
            lines=[
                LineItem("SKU-1002", "Widget B", 80, 25.00),
                LineItem("SKU-1004", "Widget D", 10, 99.50),
            ],
        ),
        Invoice(
            invoice_id="INV-201",
            vendor="Beta Parts",
            invoice_date="2024-07-01",
            currency="USD",
            layout="table_top",
            column_order="price_then_qty",
            lines=[
                LineItem("SKU-2001", "Bracket Kit", 200, 15.00),
            ],
        ),
        Invoice(
            invoice_id="INV-202",
            vendor="Beta Parts",
            invoice_date="2024-08-15",
            currency="USD",
            layout="table_middle",
            column_order="qty_then_price",
            lines=[
                LineItem("SKU-2002", "Fastener Pack", 40, 50.00),
            ],
        ),
        Invoice(
            invoice_id="INV-203",
            vendor="Beta Parts",
            invoice_date="2024-10-05",
            currency="USD",
            layout="table_bottom",
            column_order="qty_then_price",
            lines=[
                LineItem("SKU-2001", "Bracket Kit", 10, 15.00),
            ],
        ),
        Invoice(
            invoice_id="INV-301",
            vendor="Gamma Logistics",
            invoice_date="2024-08-01",
            currency="USD",
            layout="table_top",
            column_order="qty_then_price",
            lines=[
                LineItem("SKU-3001", "Shipper Carton", 100, 8.00),
            ],
        ),
        Invoice(
            invoice_id="INV-302",
            vendor="Gamma Logistics",
            invoice_date="2024-09-20",
            currency="USD",
            layout="table_middle",
            column_order="price_then_qty",
            lines=[
                LineItem("SKU-3002", "Pallet Wrap", 25, 40.00),
            ],
        ),
    ]

    # received_qty matches invoice qty except intentional mismatches below
    reports = [
        MonthlyReport(
            period="2024-07",
            rows=[
                ReportRow("SKU-1001", "Alpha Supplies", sold_qty=90, received_qty=100, return_rate=0.02),
                ReportRow("SKU-1002", "Alpha Supplies", sold_qty=40, received_qty=50, return_rate=0.01),
                ReportRow("SKU-2001", "Beta Parts", sold_qty=150, received_qty=180, return_rate=0.03),  # mismatch vs INV-201
            ],
        ),
        MonthlyReport(
            period="2024-08",
            rows=[
                ReportRow("SKU-1001", "Alpha Supplies", sold_qty=420, received_qty=450, return_rate=0.025),  # mismatch vs INV-102
                ReportRow("SKU-1003", "Alpha Supplies", sold_qty=18, received_qty=20, return_rate=0.0),
                ReportRow("SKU-2002", "Beta Parts", sold_qty=35, received_qty=40, return_rate=0.01),
                ReportRow("SKU-3001", "Gamma Logistics", sold_qty=88, received_qty=95, return_rate=0.02),  # mismatch vs INV-301
            ],
        ),
        MonthlyReport(
            period="2024-09",
            rows=[
                ReportRow("SKU-1002", "Alpha Supplies", sold_qty=70, received_qty=80, return_rate=0.015),
                ReportRow("SKU-1004", "Alpha Supplies", sold_qty=9, received_qty=10, return_rate=0.0),
                ReportRow("SKU-3002", "Gamma Logistics", sold_qty=22, received_qty=25, return_rate=0.01),
            ],
        ),
    ]

    discrepancies = [
        Discrepancy(
            id="disc-sku-1001",
            sku="SKU-1001",
            invoice_id="INV-102",
            invoice_qty=500,
            report_period="2024-08",
            report_received_qty=450,
        ),
        Discrepancy(
            id="disc-sku-2001",
            sku="SKU-2001",
            invoice_id="INV-201",
            invoice_qty=200,
            report_period="2024-07",
            report_received_qty=180,
        ),
        Discrepancy(
            id="disc-sku-3001",
            sku="SKU-3001",
            invoice_id="INV-301",
            invoice_qty=100,
            report_period="2024-08",
            report_received_qty=95,
        ),
    ]

    contract = Contract(
        vendor="Alpha Supplies",
        payment_terms="Net-30",
        effective_date="2024-01-01",
        prices=[
            ContractPrice("SKU-1001", 10.00),
            ContractPrice("SKU-1002", 25.00),
            ContractPrice("SKU-1003", 21.511),
            ContractPrice("SKU-1004", 99.50),
        ],
    )

    return Corpus(
        vendors=["Alpha Supplies", "Beta Parts", "Gamma Logistics"],
        invoices=invoices,
        reports=reports,
        discrepancies=discrepancies,
        contract=contract,
        notes=[
            "All totals are derived from line qty * unit_price.",
            "Three intentional quantity mismatches are listed under discrepancies.",
            "INV-203 is Q4 2024 and must be excluded from Alpha/Beta Q3 spend demos.",
        ],
    )


def alpha_q3_spend(corpus: Corpus) -> float:
    total = 0.0
    for inv in corpus.invoices:
        if inv.vendor == "Alpha Supplies" and inv.quarter == "2024-Q3":
            total += inv.total_amount
    return round(total, 2)


def corpus_to_ground_truth(corpus: Corpus) -> dict:
    invoices_out = []
    for inv in corpus.invoices:
        invoices_out.append(
            {
                "invoice_id": inv.invoice_id,
                "vendor": inv.vendor,
                "invoice_date": inv.invoice_date,
                "quarter": inv.quarter,
                "currency": inv.currency,
                "layout": inv.layout,
                "column_order": inv.column_order,
                "file_name": inv.file_name,
                "total_amount": inv.total_amount,
                "lines": [
                    {
                        "sku": line.sku,
                        "description": line.description,
                        "qty": line.qty,
                        "unit_price": line.unit_price,
                        "line_total": line.line_total,
                    }
                    for line in inv.lines
                ],
            }
        )

    reports_out = []
    for report in corpus.reports:
        reports_out.append(
            {
                "period": report.period,
                "file_name": report.file_name,
                "rows": [asdict(row) for row in report.rows],
            }
        )

    return {
        "version": 1,
        "currency": "USD",
        "vendors": corpus.vendors,
        "aggregates": {
            "alpha_supplies_2024_q3_spend": alpha_q3_spend(corpus),
        },
        "invoices": invoices_out,
        "reports": reports_out,
        "discrepancies": [asdict(d) for d in corpus.discrepancies],
        "contract": {
            "vendor": corpus.contract.vendor,
            "payment_terms": corpus.contract.payment_terms,
            "effective_date": corpus.contract.effective_date,
            "file_name": corpus.contract.file_name,
            "prices": [asdict(p) for p in corpus.contract.prices],
        },
        "notes": corpus.notes,
    }


def _money(value: float) -> str:
    return f"${value:,.2f}"


def _format_unit_price(value: float) -> str:
    text = f"{value:.3f}".rstrip("0").rstrip(".")
    if "." not in text:
        return f"{value:.2f}"
    return text


def write_invoice_pdf(invoice: Invoice, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
    )
    styles = getSampleStyleSheet()
    story: list = []

    header = Paragraph(f"<b>INVOICE {invoice.invoice_id}</b>", styles["Title"])
    meta = Paragraph(
        f"<b>Vendor:</b> {invoice.vendor}<br/>"
        f"<b>Date:</b> {invoice.invoice_date}<br/>"
        f"<b>Currency:</b> {invoice.currency}<br/>"
        f"<b>Bill To:</b> Demo Commerce Co.",
        styles["Normal"],
    )
    intro = Paragraph(
        "Thank you for your business. Please remit payment according to the "
        "vendor contract on file. Line items below reflect shipped units.",
        styles["Normal"],
    )
    notes = Paragraph(
        f"<i>Layout variant: {invoice.layout} / columns: {invoice.column_order}. "
        "Generated for AI-FinOps-RAG table-extraction demos.</i>",
        styles["Normal"],
    )
    total_para = Paragraph(
        f"<b>Invoice Total: {_money(invoice.total_amount)} {invoice.currency}</b>",
        styles["Heading2"],
    )

    if invoice.column_order == "price_then_qty":
        headers = ["SKU", "Description", "Unit Price", "Qty", "Line Total"]
        data = [headers]
        for line in invoice.lines:
            data.append(
                [
                    line.sku,
                    line.description,
                    _format_unit_price(line.unit_price),
                    str(line.qty),
                    f"{line.line_total:.2f}",
                ]
            )
    else:
        headers = ["SKU", "Description", "Qty", "Unit Price", "Line Total"]
        data = [headers]
        for line in invoice.lines:
            data.append(
                [
                    line.sku,
                    line.description,
                    str(line.qty),
                    _format_unit_price(line.unit_price),
                    f"{line.line_total:.2f}",
                ]
            )

    table = Table(data, hAlign="LEFT", colWidths=[1.1 * inch, 2.2 * inch, 1.0 * inch, 1.0 * inch, 1.1 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
                ("ALIGN", (2, 1), (-1, -1), "RIGHT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )

    if invoice.layout == "table_top":
        story.extend([header, Spacer(1, 8), meta, Spacer(1, 12), table, Spacer(1, 12), total_para, Spacer(1, 18), notes])
    elif invoice.layout == "table_middle":
        story.extend(
            [
                header,
                Spacer(1, 8),
                meta,
                Spacer(1, 10),
                intro,
                Spacer(1, 14),
                table,
                Spacer(1, 12),
                total_para,
                Spacer(1, 16),
                notes,
            ]
        )
    else:  # table_bottom — push table down with filler copy
        filler = Paragraph(
            "Receiving notes: Goods were inspected at dock door B. "
            "Any shortage claims must be filed within five business days. "
            "Reference the purchase order on your remittance advice. "
            "This block of text intentionally precedes the line-item table so "
            "extraction pipelines must locate tables near the bottom of the page.",
            styles["Normal"],
        )
        story.extend(
            [
                header,
                Spacer(1, 8),
                meta,
                Spacer(1, 14),
                intro,
                Spacer(1, 12),
                filler,
                Spacer(1, 20),
                filler,
                Spacer(1, 24),
                table,
                Spacer(1, 12),
                total_para,
                Spacer(1, 10),
                notes,
            ]
        )

    doc.build(story)


def write_report_xlsx(report: MonthlyReport, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "SKU Performance"
    ws.append(
        [
            "period",
            "sku",
            "vendor",
            "sold_qty",
            "received_qty",
            "return_rate",
        ]
    )
    for row in report.rows:
        ws.append(
            [
                report.period,
                row.sku,
                row.vendor,
                row.sold_qty,
                row.received_qty,
                row.return_rate,
            ]
        )
    wb.save(path)


def write_contract_docx(contract: Contract, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(f"Vendor Agreement — {contract.vendor}", level=1)
    document.add_paragraph(f"Effective date: {contract.effective_date}")
    document.add_paragraph(
        f"Payment terms: {contract.payment_terms}. Invoices are due within "
        f"{contract.payment_terms.replace('Net-', '')} days of the invoice date."
    )
    document.add_heading("Agreed unit prices (USD)", level=2)
    table = document.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "SKU"
    hdr[1].text = "Unit price"
    for price in contract.prices:
        cells = table.add_row().cells
        cells[0].text = price.sku
        cells[1].text = _format_unit_price(price.unit_price)
    document.add_paragraph(
        "Unit prices on invoices must match this schedule. A variance greater "
        "than 5% is grounds for rejection pending renegotiation."
    )
    document.save(path)


def write_golden_qa(corpus: Corpus) -> None:
    EVALS_DIR.mkdir(parents=True, exist_ok=True)
    alpha_total = alpha_q3_spend(corpus)
    beta_total = round(
        sum(
            inv.total_amount
            for inv in corpus.invoices
            if inv.vendor == "Beta Parts" and inv.quarter == "2024-Q3"
        ),
        2,
    )
    payload = {
        "description": "Golden Q&A for FinOps RAG eval — aligned with fixtures/ground_truth.json",
        "tolerance": 0.01,
        "cases": [
            {
                "id": "spend_vendor_q3",
                "question": "How much did we spend on Vendor Alpha in Q3?",
                "expect": {
                    "type": "numeric",
                    "vendor": "Alpha Supplies",
                    "period": "2024-Q3",
                    "expected_amount": alpha_total,
                    "currency": "USD",
                    "invoice_ids": ["INV-101", "INV-102", "INV-103"],
                },
            },
            {
                "id": "spend_beta_q3",
                "question": "How much did we spend on Beta Parts in Q3 2024?",
                "expect": {
                    "type": "numeric",
                    "vendor": "Beta Parts",
                    "period": "2024-Q3",
                    "expected_amount": beta_total,
                    "currency": "USD",
                    "invoice_ids": ["INV-201", "INV-202"],
                },
            },
            {
                "id": "qty_discrepancy",
                "question": "Are there quantity mismatches between invoices and product reports for SKU-1001?",
                "expect": {
                    "type": "discrepancy_alert",
                    "sku": "SKU-1001",
                    "invoice_id": "INV-102",
                    "invoice_qty": 500,
                    "report_received_qty": 450,
                    "report_period": "2024-08",
                },
            },
            {
                "id": "qty_discrepancy_beta",
                "question": "Did we receive everything billed on INV-201?",
                "expect": {
                    "type": "discrepancy_alert",
                    "sku": "SKU-2001",
                    "invoice_id": "INV-201",
                    "invoice_qty": 200,
                    "report_received_qty": 180,
                    "report_period": "2024-07",
                },
            },
            {
                "id": "qty_discrepancy_gamma",
                "question": "Are there quantity mismatches for SKU-3001?",
                "expect": {
                    "type": "discrepancy_alert",
                    "sku": "SKU-3001",
                    "invoice_id": "INV-301",
                    "invoice_qty": 100,
                    "report_received_qty": 95,
                    "report_period": "2024-08",
                },
            },
            {
                "id": "invoices_over_5000",
                "question": "Which invoices are over $5,000?",
                "expect": {
                    "type": "citation",
                    "invoice_ids": ["INV-102"],
                },
            },
            {
                "id": "payment_terms_alpha",
                "question": "What are the payment terms for Alpha Supplies?",
                "expect": {
                    "type": "citation",
                    "payment_terms": "Net-30",
                    "source_file": corpus.contract.file_name,
                },
            },
            {
                "id": "multi_discrepancy_scan",
                "question": "Are there quantity mismatches between invoices and product reports?",
                "expect": {
                    "type": "alert_count",
                    "min_count": len(corpus.discrepancies),
                },
            },
        ],
    }
    (EVALS_DIR / "golden_qa.json").write_text(
        json.dumps(payload, indent=2) + "\n",
        encoding="utf-8",
    )


def generate_all() -> dict:
    corpus = build_corpus()
    truth = corpus_to_ground_truth(corpus)

    INVOICES_DIR.mkdir(parents=True, exist_ok=True)
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    CONTRACTS_DIR.mkdir(parents=True, exist_ok=True)

    for inv in corpus.invoices:
        write_invoice_pdf(inv, INVOICES_DIR / inv.file_name)
    for report in corpus.reports:
        write_report_xlsx(report, REPORTS_DIR / report.file_name)
    write_contract_docx(corpus.contract, CONTRACTS_DIR / corpus.contract.file_name)

    (FIXTURES / "ground_truth.json").write_text(
        json.dumps(truth, indent=2) + "\n",
        encoding="utf-8",
    )
    write_golden_qa(corpus)
    return truth


def main() -> int:
    truth = generate_all()
    print(f"Wrote {len(truth['invoices'])} invoices, {len(truth['reports'])} reports, 1 contract")
    print(f"Alpha Supplies 2024-Q3 spend: ${truth['aggregates']['alpha_supplies_2024_q3_spend']:,.2f}")
    print(f"Discrepancies: {len(truth['discrepancies'])}")
    print(f"Output: {FIXTURES}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
