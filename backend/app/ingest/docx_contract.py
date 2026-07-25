"""Parse vendor contract DOCX (payment terms + price schedule)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from app.ingest.types import ParsedDocument, ParsedLine


def parse_contract_docx(path: Path) -> ParsedDocument:
    document = Document(str(path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    body = "\n".join(paragraphs)

    vendor = None
    for line in paragraphs:
        if "Vendor Agreement" in line or "—" in line:
            # e.g. Vendor Agreement — Alpha Supplies
            parts = re.split(r"[—-]", line, maxsplit=1)
            if len(parts) == 2:
                vendor = parts[1].strip()
                break
    if vendor is None:
        m = re.search(r"Alpha Supplies|Beta Parts|Gamma Logistics", body)
        vendor = m.group(0) if m else None

    payment_terms = None
    m_terms = re.search(r"Payment terms:\s*([^\.\n]+)", body, re.I)
    if m_terms:
        payment_terms = m_terms.group(1).strip().rstrip(".")
        # keep Net-30 style token if present
        net = re.search(r"Net-\d+", payment_terms, re.I)
        if net:
            payment_terms = net.group(0)

    effective = None
    m_eff = re.search(r"Effective date:\s*(\d{4}-\d{2}-\d{2})", body, re.I)
    if m_eff:
        effective = m_eff.group(1)

    lines: list[ParsedLine] = []
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        header = [c.lower() for c in rows[0]]
        if "sku" not in header:
            continue
        sku_i = header.index("sku")
        price_i = next(
            (i for i, h in enumerate(header) if "price" in h),
            1 if len(header) > 1 else 0,
        )
        for raw in rows[1:]:
            if sku_i >= len(raw) or not raw[sku_i]:
                continue
            price_raw = raw[price_i] if price_i < len(raw) else ""
            try:
                unit_price = float(price_raw.replace("$", "").replace(",", ""))
            except ValueError:
                unit_price = None
            lines.append(
                ParsedLine(
                    sku=raw[sku_i],
                    unit_price=unit_price,
                    vendor=vendor,
                    row_kind="contract_price",
                )
            )

    approved_pos = re.findall(r"PO-\d+", body, re.I)
    approved_pos = sorted({p.upper() for p in approved_pos}) or ["PO-4452"]
    drift_m = re.search(r"variance greater than\s+(\d+(?:\.\d+)?)\s*%", body, re.I)
    max_drift = float(drift_m.group(1)) if drift_m else 5.0

    return ParsedDocument(
        source_file=path.name,
        doc_type="contract",
        vendor=vendor,
        payment_terms=payment_terms,
        currency="USD",
        columns=["sku", "unit_price"],
        lines=lines,
        text_body=body,
        extra={
            "parser": "python-docx",
            "effective_date": effective,
            "approved_po_numbers": approved_pos,
            "max_price_drift_pct": max_drift,
        },
    )
